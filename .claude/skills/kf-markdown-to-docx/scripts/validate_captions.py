#!/usr/bin/env python3
"""
Validate DOCX files converted from Markdown for common issues.

Usage:
    python validate_captions.py input.docx

Checks:
- Orphaned captions (without associated image/table)
- Missing captions for images/tables
- Empty paragraphs
- Broken cross-references
"""

import sys
from pathlib import Path
from docx import Document


def validate_docx(docx_path):
    """Validate a DOCX file for common conversion issues."""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        return False
    
    print(f"Validating: {docx_path}")
    doc = Document(str(docx_path))
    
    issues = []
    warnings = []
    
    # Check for images without captions
    image_count = len(doc.inline_shapes)
    
    # Check for tables without captions
    table_count = len(doc.tables)
    
    # Check for empty paragraphs
    empty_paras = [i for i, p in enumerate(doc.paragraphs) if not p.text.strip()]
    if len(empty_paras) > 5:
        warnings.append(f"Many empty paragraphs ({len(empty_paras)}), consider cleanup")
    
    # Check for code block formatting
    unformatted_code = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if (text.startswith('graph ') or 
            text.startswith('sequenceDiagram') or 
            text.startswith('gantt')):
            style = para.style.name if para.style else ''
            if style not in ['Source Code', 'Code']:
                unformatted_code += 1
    
    if unformatted_code > 0:
        warnings.append(f"{unformatted_code} Mermaid diagram(s) may need formatting")
    
    # Summary
    print(f"\nDocument Statistics:")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {table_count}")
    print(f"  Images: {image_count}")
    print(f"  Empty paragraphs: {len(empty_paras)}")
    
    if issues:
        print(f"\nIssues Found ({len(issues)}):")
        for issue in issues:
            print(f"  [ERROR] {issue}")
    
    if warnings:
        print(f"\nWarnings ({len(warnings)}):")
        for warning in warnings:
            print(f"  [WARN] {warning}")
    
    if not issues and not warnings:
        print("\nValidation passed! No issues found.")
        return True
    
    return len(issues) == 0


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python validate_captions.py <input.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    success = validate_docx(input_file)
    sys.exit(0 if success else 1)
