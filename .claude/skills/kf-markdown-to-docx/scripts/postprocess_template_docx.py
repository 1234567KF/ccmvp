#!/usr/bin/env python3
"""
Post-process DOCX files converted from Markdown to fix common formatting issues.

Usage:
    python postprocess_template_docx.py input.docx

Fixes applied:
- Caption spacing
- Table borders
- Code block styling (monospace font, background)
- List indentation
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_code_paragraph(para):
    """Format a code block paragraph with monospace font and background."""
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    # Add light gray background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    para.paragraph_format.element.get_or_add_pPr().append(shading_elm)
    
    # Add spacing
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def fix_table_borders(table):
    """Ensure tables have visible borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def apply_table_text_style(table, doc):
    """Apply '表格文本' style to all table cells if available in template."""
    # Check if template has '表格文本' style
    style_names = [s.name for s in doc.styles]
    has_table_text_style = '表格文本' in style_names
    
    if not has_table_text_style:
        return 0
    
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.style = doc.styles['表格文本']
                count += 1
    return count


def process_docx(docx_path):
    """Process a DOCX file with post-processing fixes."""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        return False
    
    print(f"Processing: {docx_path}")
    doc = Document(str(docx_path))
    
    # Track modifications
    code_blocks_formatted = 0
    tables_fixed = 0
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect code blocks (paragraphs with Source Code style or containing code-like content)
        style_name = para.style.name if para.style else ''
        if style_name == 'Source Code' or style_name == 'Code':
            format_code_paragraph(para)
            code_blocks_formatted += 1
        
        # Detect and format Mermaid diagram text (starts with graph, sequenceDiagram, etc.)
        if (text.startswith('graph ') or 
            text.startswith('sequenceDiagram') or 
            text.startswith('gantt') or
            text.startswith('flowchart')):
            format_code_paragraph(para)
            code_blocks_formatted += 1
    
    # Process tables
    table_cells_styled = 0
    for table in doc.tables:
        fix_table_borders(table)
        table_cells_styled += apply_table_text_style(table, doc)
        tables_fixed += 1
    
    # Save
    doc.save(str(docx_path))
    
    print(f"Done!")
    print(f"  Code blocks formatted: {code_blocks_formatted}")
    print(f"  Tables fixed: {tables_fixed}")
    print(f"  Table cells styled: {table_cells_styled}")
    
    return True


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python postprocess_template_docx.py <input.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    success = process_docx(input_file)
    sys.exit(0 if success else 1)
