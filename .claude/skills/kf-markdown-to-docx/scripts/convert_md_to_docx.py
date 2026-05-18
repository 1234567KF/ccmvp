#!/usr/bin/env python3
"""
Convert Markdown to DOCX with comprehensive formatting:
1. Mermaid diagrams rendered to PNG (mermaid.ink API primary, local CLI fallback)
2. Word template styles applied (Heading 1-4, 琛ㄦ牸鏂囨湰, etc.)
3. Table borders and formatting
4. Code block styling

Usage:
    python convert_md_to_docx.py input.md template.docx output.docx
"""

import re
import os
import sys
import base64
import hashlib
import shutil
import subprocess
import requests
from pathlib import Path
# Auto-install dependencies on first run
def _ensure_dependencies():
    missing = []
    try:
        __import__('docx')
    except ImportError:
        missing.append('python-docx')
    try:
        __import__('pypandoc')
    except ImportError:
        missing.append('pypandoc')
    try:
        __import__('requests')
    except ImportError:
        missing.append('requests')
    if missing:
        print(f'[Auto-install] Missing packages: {missing}')
        setup_script = __import__('pathlib').Path(__file__).parent / 'setup.py'
        if setup_script.exists():
            print(f'[Auto-install] Running setup.py...')
            result = __import__('subprocess').run([__import__('sys').executable, str(setup_script)], capture_output=True, text=True)
            if result.returncode != 0:
                print(f'[Auto-install] Setup failed. Please run manually.')
                __import__('sys').exit(1)
        else:
            print(f'[Auto-install] Please install: pip install ' + ' '.join(missing))
            __import__('sys').exit(1)
_ensure_dependencies()

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pypandoc


def render_mermaid_via_api(mermaid_code, output_path):
    """Render mermaid diagram to PNG using mermaid.ink API."""
    try:
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            output_path.write_bytes(response.content)
            return True
    except Exception:
        pass
    return False


def render_mermaid_via_cli(mermaid_code, output_path, cwd=None):
    """Render mermaid diagram using local mermaid-cli with system Chrome."""
    temp_mmd = None
    try:
        temp_mmd = output_path.with_suffix('.mmd')
        temp_mmd.write_text(mermaid_code, encoding='utf-8')
        
        # Detect system Chrome path
        chrome_paths = [
            'C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe',
            'C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe',
        ]
        local_app_data = os.environ.get('LOCALAPPDATA', '')
        if local_app_data:
            chrome_paths.append(
                os.path.join(local_app_data, 'Google', 'Chrome', 'Application', 'chrome.exe')
            )
        chrome_path = None
        for cp in chrome_paths:
            if os.path.exists(cp):
                chrome_path = cp
                break
        
        env = os.environ.copy()
        if chrome_path:
            env['PUPPETEER_EXECUTABLE_PATH'] = chrome_path
        
        # On Windows, npx is a .CMD file; use shell=True for reliable execution
        is_windows = sys.platform == 'win32'
        npx_path = shutil.which('npx') or 'npx'
        cmd = [npx_path, 'mmdc', '-i', str(temp_mmd), '-o', str(output_path)]
        
        result = subprocess.run(
            cmd,
            check=True, capture_output=True, timeout=120, env=env,
            shell=is_windows, cwd=cwd
        )
        if temp_mmd.exists():
            temp_mmd.unlink(missing_ok=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
        # Print error details for debugging
        print(f"    CLI error: {type(e).__name__}")
        if hasattr(e, 'stderr') and e.stderr:
            err_text = e.stderr.decode('utf-8', errors='ignore')[:300]
            print(f"    {err_text}")
        if temp_mmd and temp_mmd.exists():
            temp_mmd.unlink(missing_ok=True)
        return False


def render_mermaid(mermaid_code, output_path, cwd=None):
    """Render mermaid diagram with fallback mechanisms."""
    # Try local CLI first (more reliable for complex diagrams)
    if render_mermaid_via_cli(mermaid_code, output_path, cwd=cwd):
        return 'cli'
    # Fallback to API
    if render_mermaid_via_api(mermaid_code, output_path):
        return 'api'
    return None


def process_mermaid_diagrams(content, images_dir, cwd=None):
    """Find and render all mermaid diagrams in markdown content."""
    mermaid_pattern = r'```mermaid\n(.*?)\n```'
    matches = list(re.finditer(mermaid_pattern, content, re.DOTALL))
    
    if not matches:
        return content, 0, 0
    
    print(f"Found {len(matches)} Mermaid diagram(s)")
    rendered = 0
    failed = 0
    replacements = []
    
    for i, match in enumerate(matches):
        mermaid_code = match.group(1).strip()
        content_hash = hashlib.md5(mermaid_code.encode()).hexdigest()[:8]
        image_name = f'diagram_{i+1}_{content_hash}.png'
        image_path = images_dir / image_name
        
        print(f"Rendering diagram {i+1}/{len(matches)}...")
        method = render_mermaid(mermaid_code, image_path, cwd=cwd)
        
        if method:
            # Use absolute path to ensure pandoc can find images regardless of cwd
            abs_path = image_path.resolve()
            replacement = f'![Diagram {i+1}]({abs_path})'
            replacements.append((match.start(), match.end(), replacement))
            print(f"  OK ({method})")
            rendered += 1
        else:
            print(f"  Failed - keeping as code")
            failed += 1
    
    # Apply replacements in reverse order
    new_content = content
    for start, end, replacement in reversed(replacements):
        new_content = new_content[:start] + replacement + new_content[end:]
    
    return new_content, rendered, failed


def fix_table_borders(table):
    """Ensure tables have visible borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def apply_table_styles(doc):
    """Apply template table text style to all tables if available."""
    style_names = [s.name for s in doc.styles]
    has_table_style = '琛ㄦ牸鏂囨湰' in style_names
    
    if not has_table_style:
        return 0
    
    count = 0
    for table in doc.tables:
        fix_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.style = doc.styles['琛ㄦ牸鏂囨湰']
                    count += 1
    return count


def format_code_blocks(doc):
    """Format code block paragraphs with monospace font and background."""
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ''
        
        is_code = (style_name in ['Source Code', 'Code'] or
                   text.startswith('graph ') or 
                   text.startswith('sequenceDiagram') or 
                   text.startswith('gantt') or
                   text.startswith('flowchart'))
        
        if is_code:
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            para.paragraph_format.element.get_or_add_pPr().append(shading_elm)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            count += 1
    
    return count


def convert_md_to_docx(md_path, template_path, output_path):
    """Main conversion function."""
    md_path = Path(md_path)
    template_path = Path(template_path)
    output_path = Path(output_path)
    
    if not md_path.exists():
        print(f"Error: Markdown file not found: {md_path}")
        return False
    
    if not template_path.exists():
        print(f"Error: Template file not found: {template_path}")
        return False
    
    # Create images directory
    images_dir = md_path.parent / 'mermaid-images'
    images_dir.mkdir(exist_ok=True)
    
    # Read and process markdown
    content = md_path.read_text(encoding='utf-8')
    new_content, rendered, failed = process_mermaid_diagrams(content, images_dir, cwd=str(md_path.parent))
    
    print(f"\nDiagrams: {rendered} rendered, {failed} kept as code")
    
    # Write temporary markdown
    temp_md = md_path.parent / f'.temp_{md_path.stem}.md'
    temp_md.write_text(new_content, encoding='utf-8')
    
    # Convert to docx
    print(f"\nConverting to DOCX...")
    try:
        pypandoc.convert_file(
            str(temp_md),
            'docx',
            outputfile=str(output_path),
            extra_args=[f'--reference-doc={template_path}']
        )
    except Exception as e:
        print(f"Template conversion failed: {e}")
        print("Falling back to basic conversion...")
        pypandoc.convert_file(str(temp_md), 'docx', outputfile=str(output_path))
    finally:
        temp_md.unlink(missing_ok=True)
    
    # Post-process
    print(f"Post-processing...")
    doc = Document(str(output_path))
    
    table_cells = apply_table_styles(doc)
    code_blocks = format_code_blocks(doc)
    
    doc.save(str(output_path))
    
    print(f"\nDone! Output: {output_path}")
    print(f"  Images embedded: {rendered}")
    print(f"  Table cells styled: {table_cells}")
    print(f"  Code blocks formatted: {code_blocks}")
    
    return True


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python convert_md_to_docx.py <input.md> <template.docx> <output.docx>")
        sys.exit(1)
    
    success = convert_md_to_docx(sys.argv[1], sys.argv[2], sys.argv[3])
    sys.exit(0 if success else 1)
